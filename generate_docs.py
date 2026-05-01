"""
generate_docs.py — Generate Nexora_AI_Platform_Documentation.docx
Run: python generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

DOC_NAME = "Nexora_AI_Platform_Documentation.docx"
PURPLE   = RGBColor(0x7C, 0x3A, 0xED)
ACCENT   = RGBColor(0xBC, 0x8C, 0xFF)
DARK     = RGBColor(0x1F, 0x2D, 0x3D)
MUTED    = RGBColor(0x6B, 0x72, 0x80)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p


def add_body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold  = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p


def add_bullet(doc, items, indent=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if ":" in item and item.index(":") < 40:
            parts = item.split(":", 1)
            r1 = p.add_run(parts[0] + ":")
            r1.bold = True
            r1.font.color.rgb = DARK
            r2 = p.add_run(parts[1])
            r2.font.color.rgb = MUTED
        else:
            run = p.add_run(item)
            run.font.color.rgb = DARK
        p.paragraph_format.left_indent = Cm(indent)


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.color.rgb = DARK


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x18, 0x18, 0x18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shading_elm = OxmlElement("w:pPr")
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cell, "7C3AED")

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_bg(cell, "F5F3FF")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BC8CFF")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("⚡ NEXORA")
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = PURPLE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run("AI PLATFORM")
    sub_r.bold = True
    sub_r.font.size = Pt(22)
    sub_r.font.color.rgb = ACCENT

    doc.add_paragraph("")
    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_p.add_run("Complete Technical Documentation").font.size = Pt(14)

    doc.add_paragraph("")
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_r = date_p.add_run(f"Version 2.0  ·  {datetime.date.today().strftime('%B %Y')}")
    date_r.font.color.rgb = MUTED
    date_r.font.size = Pt(11)

    doc.add_paragraph("")
    doc.add_paragraph("")
    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_r = conf_p.add_run("CONFIDENTIAL — INVESTOR / TECHNICAL EDITION")
    conf_r.font.color.rgb = MUTED
    conf_r.font.size = Pt(9)
    conf_r.italic = True

    doc.add_page_break()

    # ── 1. PROJECT OVERVIEW ──────────────────────────────────────────────
    add_heading(doc, "1. Project Overview", 1, PURPLE)
    add_divider(doc)
    add_body(doc, "Nexora is a production-grade, multi-agent AI development platform that enables developers to build software autonomously. Users describe what they want — and the AI plans, codes, debugs, and delivers working software step by step.")
    doc.add_paragraph("")
    add_body(doc, "The platform operates on a continuous Plan → Execute → Observe → Fix loop, with 36 distinct capability phases built over the core engine. It supports 15+ AI providers, 5 specialist coding agents, a 3-tier memory system, autonomous goal execution, background task scheduling, and a full SaaS billing system powered by Razorpay.", size=10)
    doc.add_paragraph("")
    add_body(doc, "Key Platform Metrics:", bold=True)
    add_bullet(doc, [
        "~10,000 lines of backend Python (web_app.py)",
        "~16,000 lines of frontend JavaScript/HTML (index.html)",
        "36 feature phases, all production-tested",
        "15+ AI provider integrations",
        "2 SQLite databases (sessions.db, billing.db)",
        "Full SaaS billing with Razorpay",
    ])
    doc.add_paragraph("")

    # ── 2. SYSTEM ARCHITECTURE ──────────────────────────────────────────
    add_heading(doc, "2. System Architecture", 1, PURPLE)
    add_divider(doc)
    add_heading(doc, "2.1 Backend Architecture", 2)
    add_bullet(doc, [
        "web_app.py: Flask application, all 36 phases, ~10,000 lines",
        "agent.py: Core autonomous agent loop (Plan-Execute-Observe-Fix)",
        "orchestrator.py: Multi-agent pipeline orchestration",
        "router.py / model_router.py: Intelligent LLM selection and routing",
        "memory.py: 3-tier context management system",
        "payments.py: Full Razorpay billing integration",
        "auth_system.py: JWT-based authentication",
        "tools.py + tools/: Extensible tool registry",
        "vector_store.py + long_term_memory.py: Semantic memory via ChromaDB",
    ])
    add_heading(doc, "2.2 Frontend Architecture", 2)
    add_bullet(doc, [
        "Single-page application in templates/index.html (~16,000 lines)",
        "Three-panel responsive layout: AI Thinking / Logs-Code-Terminal / Inspector",
        "Dark/light theme with CSS custom properties",
        "Real-time polling for agent status, queue, and session updates",
        "All 36 phases integrated into the single SPA",
    ])
    add_heading(doc, "2.3 Data Layer", 2)
    add_table(doc,
        ["Database", "Tables", "Purpose"],
        [
            ["sessions.db", "sessions, settings, scheduler, chat_summaries, errors", "Core app state"],
            ["billing.db", "subscriptions, invoices, payment_events", "SaaS billing"],
        ],
        col_widths=[1.5, 3.0, 2.0]
    )
    doc.add_paragraph("")

    # ── 3. FEATURE BREAKDOWN ─────────────────────────────────────────────
    add_heading(doc, "3. Feature Breakdown (Phase 1 → Phase 36)", 1, PURPLE)
    add_divider(doc)
    add_table(doc,
        ["Phase", "Feature", "Description"],
        [
            ["1–3", "Core Agent Loop", "Plan → Execute → Observe → Fix autonomous development loop"],
            ["4", "Intelligence & Personalization", "Adaptive responses, session learning, preference tracking"],
            ["5", "BYOK Multi-Provider", "15+ AI providers with cost-aware routing and auto-failover"],
            ["6", "Decision Intelligence", "Priority lock, fast/smart/cheap routing modes"],
            ["7", "Structured Agents", "5 specialist agents triggered by task context"],
            ["8", "Monetization Layer", "Plan gating (Free/Pro/Elite) with server-side usage quotas"],
            ["9", "Model Intelligence Routing", "Per-role model assignment (planning, coding, debug)"],
            ["10", "Agent Intelligence", "Short/long-term memory, self-correction from past errors"],
            ["11", "Multi-Agent Collaboration", "4-agent pipeline: Manager, Research, Coding, Debug"],
            ["12", "Editable Prompt System", "Inline prompt editing, conversation branching"],
            ["13", "3-tier Context Compression", "Intelligent summarisation to maximise context window"],
            ["14", "Self-Improving AI", "Post-session reflection and capability improvement"],
            ["15", "Learning Dashboard", "Success metrics, failure analysis, timeline"],
            ["16", "Autonomous Goal AI", "Goal decomposition into executable task graphs"],
            ["17", "Task Graph Visualiser", "DAG visualisation of execution graphs"],
            ["18", "Background Scheduler", "Cron-like autonomous agent scheduling"],
            ["19", "Security Hardening", "Rate limiting, JWT auth, kill switch, admin controls"],
            ["20–32", "Platform Extensions", "Terminal intelligence, cross-platform tools, cost tracking"],
            ["33", "Cross-Platform Terminal", "Full terminal bridging and execution intelligence"],
            ["34", "Semi-AGI Dashboard", "Advanced intelligence metrics and reasoning traces"],
            ["35", "Enterprise Dashboards", "Business analytics, performance and ROI reporting"],
            ["36", "SaaS Billing (Razorpay)", "Orders, verification, webhooks, invoices, email receipts"],
        ],
        col_widths=[0.7, 2.3, 3.5]
    )
    doc.add_paragraph("")

    # ── 4. AI CAPABILITIES ───────────────────────────────────────────────
    add_heading(doc, "4. AI Capabilities", 1, PURPLE)
    add_divider(doc)
    add_heading(doc, "4.1 Planning", 2)
    add_body(doc, "The planning phase uses high-reasoning models (GPT-4o, Claude 3.5, Gemini Pro) to decompose user requirements into structured, ordered sub-tasks. Each task includes context, dependencies, and expected output.")
    add_heading(doc, "4.2 Coding", 2)
    add_body(doc, "Code-optimised models (GPT-4o, Claude 3.5 Sonnet, DeepSeek Coder) implement each sub-task, with access to file system tools, shell execution, and web search. Code is written incrementally with verification at each step.")
    add_heading(doc, "4.3 Debugging", 2)
    add_body(doc, "The debug agent analyses errors, traces root causes, and applies targeted fixes. It has access to error logs, test output, and past failure patterns from the memory system.")
    add_heading(doc, "4.4 Multi-Agent System (Phase 11)", 2)
    add_body(doc, "For complex tasks, a 4-agent pipeline coordinates automatically:")
    add_numbered(doc, [
        "Manager Agent — Decomposes goals, assigns sub-tasks, tracks progress",
        "Research Agent — Gathers context, documentation, relevant examples",
        "Coding Agent — Implements the solution, writes and refines code",
        "Debug Agent — Tests, identifies bugs, ensures output correctness",
    ])
    add_heading(doc, "4.5 Goal Execution (Phase 16)", 2)
    add_body(doc, "Users define high-level goals. The system automatically decomposes them into a directed acyclic graph (DAG) of sub-tasks, executes them in dependency order, and reports progress in real time via the Task Graph Visualiser.")
    doc.add_paragraph("")

    # ── 5. MODEL ROUTING ─────────────────────────────────────────────────
    add_heading(doc, "5. Model Routing System (Phase 9)", 1, PURPLE)
    add_divider(doc)
    add_body(doc, "The model routing system intelligently assigns the best available AI model for each task role based on plan tier, user priority preference, provider availability, and cost constraints.")
    add_table(doc,
        ["Routing Mode", "Description", "Best For"],
        [
            ["Lite", "Cost-optimised, fast models (Groq, Flash)", "Quick tasks, high-volume usage"],
            ["Pro", "Balanced cost and quality", "General development tasks"],
            ["Elite", "Highest quality models (GPT-4o, Claude 3.5)", "Complex architecture, critical code"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )
    doc.add_paragraph("")
    add_body(doc, "Auto-Failover: If the primary provider rate-limits or errors, the router automatically tries the next best provider — transparently, with no user intervention.", italic=True)
    doc.add_paragraph("")

    # ── 6. MEMORY SYSTEM ─────────────────────────────────────────────────
    add_heading(doc, "6. Memory System (Phases 10, 13, 14)", 1, PURPLE)
    add_divider(doc)
    add_bullet(doc, [
        "Tier 1 — Short-term: Full in-session message history (lossless)",
        "Tier 2 — Mid-term: Compressed summaries of older context, preserving key decisions",
        "Tier 3 — Long-term: Semantic vector store via ChromaDB for cross-session recall",
    ])
    add_body(doc, "The system tracks errors and successes. Before each new task, relevant past failures are surfaced to the agent, preventing repeated mistakes. An Intelligence Score (A–F grade) tracks agent performance per session.", size=10)
    doc.add_paragraph("")

    # ── 7. SCHEDULER ─────────────────────────────────────────────────────
    add_heading(doc, "7. Scheduler System (Phase 18)", 1, PURPLE)
    add_divider(doc)
    add_body(doc, "The background task scheduler enables autonomous agent operation without user interaction.")
    add_bullet(doc, [
        "Cron-like scheduling (minute, hour, day, week intervals)",
        "Background AI agent execution",
        "Email notifications on task completion",
        "Task history, status tracking, and retry on failure",
        "Integration with the goal system for scheduled goal execution",
    ])
    doc.add_paragraph("")

    # ── 8. SECURITY ──────────────────────────────────────────────────────
    add_heading(doc, "8. Security & Hardening (Phase 19)", 1, PURPLE)
    add_divider(doc)
    add_table(doc,
        ["Security Feature", "Implementation"],
        [
            ["Authentication", "JWT tokens, bcrypt password hashing, token expiry"],
            ["Rate Limiting", "Per-endpoint limits with 429 responses"],
            ["Input Sanitisation", "All user inputs sanitised before processing"],
            ["SQL Injection Prevention", "Parameterised queries throughout"],
            ["Webhook Verification", "HMAC-SHA256 signature verification for all Razorpay webhooks"],
            ["CORS", "Configurable CORS headers on all responses"],
            ["Kill Switch", "/api/admin/kill-switch for emergency shutdown"],
            ["Structured Logging", "No secrets in logs, all errors captured with context"],
            ["WAL Mode SQLite", "Crash-safe writes to all databases"],
        ],
        col_widths=[2.5, 4.0]
    )
    doc.add_paragraph("")

    # ── 9. MONETIZATION ──────────────────────────────────────────────────
    add_heading(doc, "9. Monetization System (Phases 8 & 36)", 1, PURPLE)
    add_divider(doc)
    add_table(doc,
        ["Plan", "Monthly Price", "Daily Runs", "Key Features"],
        [
            ["Free", "₹0", "10", "Lite routing, basic agents"],
            ["Pro", "₹20", "Unlimited", "Pro routing, all 5 agents, memory, BYOK"],
            ["Elite", "₹50", "Unlimited", "Elite routing, multi-agent, scheduler, goals, priority support"],
        ],
        col_widths=[1.0, 1.5, 1.2, 2.8]
    )
    doc.add_paragraph("")
    add_heading(doc, "Razorpay Payment Flow", 2)
    add_numbered(doc, [
        "User selects plan and billing cycle (Monthly/Yearly)",
        "Frontend calls /api/payments/create-order — Razorpay order created",
        "Razorpay Checkout popup opens in browser",
        "User completes payment via card/UPI/netbanking",
        "Frontend calls /api/payments/verify with HMAC signature",
        "Subscription activated in billing.db",
        "Invoice generated and stored",
        "Email receipt sent via Resend API",
        "Razorpay webhook /api/payments/webhook provides idempotent backup activation",
    ])
    doc.add_paragraph("")

    # ── 10. BYOK ────────────────────────────────────────────────────────
    add_heading(doc, "10. BYOK System (Phase 5)", 1, PURPLE)
    add_divider(doc)
    add_body(doc, "The Bring Your Own Key (BYOK) system allows users to supply their own API keys for any supported provider. Keys are stored encrypted in the session database and used in preference to platform keys.")
    add_bullet(doc, [
        "15+ providers supported (see API Providers section)",
        "Keys stored per-user-session, never logged",
        "Automatic fallback to platform keys when user key is absent",
        "Provider health check and balance display",
        "Key visibility toggle (show/hide) in UI",
    ])
    doc.add_paragraph("")

    # ── 11. DEPLOYMENT ───────────────────────────────────────────────────
    add_heading(doc, "11. Deployment Guide", 1, PURPLE)
    add_divider(doc)
    add_heading(doc, "Replit Deployment (Recommended)", 2)
    add_numbered(doc, [
        "Import the project to Replit",
        "Add all secrets in Replit's Secrets panel (see Environment Variables section)",
        "Click Deploy — the platform automatically provisions TLS, health checks, and auto-restart",
        "Your app will be available at https://your-project.replit.app",
        "Configure Razorpay webhook URL using the production domain",
    ])
    add_heading(doc, "Self-Hosted Deployment", 2)
    add_code(doc, "pip install -r requirements.txt\ngunicorn -w 2 -b 0.0.0.0:5000 web_app:app")
    doc.add_paragraph("")
    add_heading(doc, "Production Checklist", 2)
    add_bullet(doc, [
        "SECRET_KEY and JWT_SECRET set to 32+ character random strings",
        "Razorpay live keys (not test keys) in production",
        "Webhook URL configured: https://your-domain.com/api/payments/webhook",
        "EMAIL_FROM uses a verified sender domain",
        "HTTPS enabled (automatic on Replit)",
    ])
    doc.add_paragraph("")

    # ── 12. ENVIRONMENT VARIABLES ────────────────────────────────────────
    add_heading(doc, "12. Environment Variables", 1, PURPLE)
    add_divider(doc)
    add_table(doc,
        ["Variable", "Required", "Description"],
        [
            ["SECRET_KEY", "✓ Required", "Flask session signing secret (32+ chars)"],
            ["JWT_SECRET", "✓ Required", "JWT token signing secret (32+ chars)"],
            ["RAZORPAY_KEY_ID", "Billing", "Razorpay API key ID from dashboard"],
            ["RAZORPAY_KEY_SECRET", "Billing", "Razorpay secret key from dashboard"],
            ["EMAIL_API_KEY", "Optional", "Resend.com API key for invoice emails"],
            ["EMAIL_FROM", "Optional", "Sender address (e.g. billing@nexora.ai)"],
            ["OPENAI_API_KEY", "Optional", "OpenAI API key (platform-level)"],
            ["GEMINI_API_KEY", "Optional", "Google Gemini API key"],
            ["ANTHROPIC_API_KEY", "Optional", "Anthropic Claude API key"],
            ["GROQ_API_KEY", "Optional", "Groq API key"],
        ],
        col_widths=[2.0, 1.2, 3.3]
    )
    doc.add_paragraph("")

    # ── 13. API ENDPOINTS ────────────────────────────────────────────────
    add_heading(doc, "13. API Endpoints Summary", 1, PURPLE)
    add_divider(doc)
    add_heading(doc, "Core Endpoints", 2)
    add_table(doc,
        ["Endpoint", "Method", "Description"],
        [
            ["/", "GET", "Main application SPA"],
            ["/docs", "GET", "Platform documentation"],
            ["/api/queue-task", "POST", "Submit task to agent queue"],
            ["/api/queue", "GET", "Queue status and history"],
            ["/api/sessions", "GET/POST", "Session management"],
            ["/api/providers", "GET", "Available AI providers"],
            ["/api/plan/info", "GET", "Current subscription info"],
            ["/api/health", "GET", "Health check"],
            ["/api/auth/signup", "POST", "Create user account"],
            ["/api/auth/login", "POST", "Login and get JWT"],
        ],
        col_widths=[2.5, 1.0, 3.0]
    )
    add_heading(doc, "Billing Endpoints", 2)
    add_table(doc,
        ["Endpoint", "Method", "Description"],
        [
            ["/api/payments/plans", "GET", "Available plans and pricing"],
            ["/api/payments/create-order", "POST", "Create Razorpay order"],
            ["/api/payments/verify", "POST", "Verify payment and activate plan"],
            ["/api/payments/webhook", "POST", "Razorpay webhook receiver"],
            ["/api/billing/info", "GET", "Subscription and invoice info"],
            ["/api/billing/invoices", "GET", "Invoice list"],
            ["/api/billing/cancel", "POST", "Cancel subscription"],
            ["/api/invoice/<id>", "GET", "Download invoice HTML"],
        ],
        col_widths=[2.5, 1.0, 3.0]
    )
    doc.add_paragraph("")

    # ── 14. UI WALKTHROUGH ───────────────────────────────────────────────
    add_heading(doc, "14. UI Walkthrough", 1, PURPLE)
    add_divider(doc)
    add_heading(doc, "Three-Panel Layout", 2)
    add_bullet(doc, [
        "Left Panel — AI Thinking: Live reasoning traces, session history, agent actions",
        "Centre Panel — Logs/Code/Terminal: Task output, file editor, integrated terminal, metrics",
        "Right Panel — Inspector: Model info, plan status, billing, intelligence score, memory",
    ])
    add_heading(doc, "Header Controls", 2)
    add_bullet(doc, [
        "Nexora logo and branding",
        "Plan mode selector (Lite / Pro / Elite) with routing dropdown",
        "Subscription badge showing current plan and expiry date",
        "Theme toggle (dark/light)",
        "Token and cost counter",
        "Run task button",
    ])
    add_heading(doc, "Upgrade / Billing Modal", 2)
    add_bullet(doc, [
        "Billing cycle toggle (Monthly / Yearly)",
        "Plan cards with INR pricing and feature lists",
        "Real Razorpay checkout popup on 'Pay' click",
        "Subscription status chip and expiry display",
        "Invoice history with download links",
        "Billing Setup Guide with webhook URL and connection status",
    ])
    doc.add_paragraph("")

    # ── 15. FUTURE SCOPE ─────────────────────────────────────────────────
    add_heading(doc, "15. Future Scope", 1, PURPLE)
    add_divider(doc)
    add_bullet(doc, [
        "WebSocket real-time streaming (replace polling)",
        "Team / organisation accounts with role-based access control",
        "GitHub integration: commit, PR creation, code review automation",
        "Plugin marketplace for custom agent tools",
        "Mobile application (React Native / Expo)",
        "Multi-tenant SaaS mode with per-tenant isolation",
        "Annual billing plans via Razorpay Subscriptions API",
        "Voice interface for task submission",
        "Visual code editor with live preview",
        "SOC2 compliance and audit logging",
    ])

    # ── FOOTER ──────────────────────────────────────────────────────────
    doc.add_paragraph("")
    add_divider(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run(f"⚡ Nexora AI Platform · Official Documentation · {datetime.date.today().year}")
    footer_r.font.color.rgb = MUTED
    footer_r.font.size = Pt(9)
    footer_r.italic = True

    doc.save(DOC_NAME)
    print(f"✓ Documentation saved: {DOC_NAME}")


if __name__ == "__main__":
    build_doc()
